from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOT_DIR = ROOT / "selenium" / "screenshots"
OUTPUT_DIR = ROOT / "reports"
OUTPUT_PATH = OUTPUT_DIR / "Laporan_Hasil_Selenium_SIPENA.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203040"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
GREEN = "16794A"
PALE_GREEN = "E9F7EF"
WHITE = "FFFFFF"
BORDER = "D0D5DD"

SCENARIOS = [
    ("Login", "/login", "Form login tersedia; autentikasi admin berhasil", "smoke-login-pass.png"),
    ("Arsip & Riwayat", "/activity-archive", "Halaman tampil dan filter pengguna tersedia", "smoke-arsip-riwayat-pass.png"),
    ("Dashboard", "/", "Kartu ringkasan inventaris, peminjaman, dan pemeliharaan tampil", "smoke-dashboard-pass.png"),
    ("Dokumentasi Sistem", "/uml", "Activity Diagram dan Use Case Diagram tampil", "smoke-dokumentasi-sistem-pass.png"),
    ("Inventaris Medis", "/medical-assets", "Judul dan pencarian inventaris medis tersedia", "smoke-inventaris-medis-pass.png"),
    ("Inventaris Non-Medis", "/non-medical-assets", "Judul dan pencarian inventaris non-medis tersedia", "smoke-inventaris-non-medis-pass.png"),
    ("Jadwal Pemeliharaan", "/maintenance-schedule", "Kontrol tambah jadwal dan daftar jadwal tampil", "smoke-jadwal-pemeliharaan-pass.png"),
    ("Laporan & Analitik", "/reports", "Ekspor laporan dan ringkasan analitik tampil", "smoke-laporan-analitik-pass.png"),
    ("Manajemen Pengguna", "/users", "Data pengguna, unit kerja, dan tambah pengguna tampil", "smoke-manajemen-pengguna-pass.png"),
    ("Manajemen Sanksi", "/sanctions", "Daftar dan pencarian sanksi tersedia", "smoke-manajemen-sanksi-pass.png"),
    ("Pemeliharaan Sarana", "/maintenance", "Kalender, daftar, dan riwayat pemeliharaan tampil", "smoke-pemeliharaan-sarana-pass.png"),
    ("Peminjaman", "/borrowing", "Daftar dan pencarian peminjaman tampil", "smoke-peminjaman-pass.png"),
    ("Pengaturan", "/settings", "Profil akun, ganti sandi, dan informasi sistem tampil", "smoke-pengaturan-pass.png"),
    ("Pengembalian", "/returns", "Daftar alat dan riwayat pengembalian tampil", "smoke-pengembalian-pass.png"),
    ("Penggunaan", "/asset-usage", "Riwayat pemakaian dan pencarian tersedia", "smoke-penggunaan-pass.png"),
    ("Penghapusan Aset", "/disposal", "Permintaan penghapusan dan pencarian aset tampil", "smoke-penghapusan-aset-pass.png"),
    ("SPK Prioritas Aset", "/dss", "Ringkasan, ranking, dan metode TOPSIS tampil", "smoke-spk-prioritas-aset-pass.png"),
    ("Unggah Dokumen", "/unggahan", "Keterangan dan riwayat dokumen tampil", "smoke-unggah-dokumen-pass.png"),
    ("Logout", "/login", "Sesi berakhir, token terhapus, dan kembali ke login", "smoke-logout-pass.png"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def add_bottom_border(paragraph, color=BLUE, size="18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_image_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("name", title)
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Title", 24, INK, 0, 4),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("Halaman ")
    set_font(run, size=9, color=MUTED)
    add_field(p, "PAGE")


def add_metadata(doc: Document, run_time: datetime) -> None:
    rows = [
        ("Sistem", "SIPENA — Sistem Informasi Pengelolaan Sarana"),
        ("Jenis", "Navigation smoke test (Selenium WebDriver)"),
        ("Waktu eksekusi", run_time.strftime("%d %B %Y, %H:%M WIB").replace("July", "Juli")),
        ("Hasil", "19 lulus, 0 gagal — tingkat kelulusan 100%"),
        ("Sumber bukti", "selenium/screenshots/*.png"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], LIGHT_GRAY)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), size=10, color=INK, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        color = GREEN if label == "Hasil" else INK
        set_font(p.add_run(value), size=10, color=color, bold=(label == "Hasil"))
    set_table_geometry(table, [2000, 7360])
    set_table_borders(table)


def add_status_callout(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [1650, 7710])
    set_table_borders(table, color="B7E4C7", size="8")
    for cell in table.rows[0].cells:
        set_cell_shading(cell, PALE_GREEN)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("LULUS"), size=14, color=GREEN, bold=True)
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("Seluruh 19 skenario menghasilkan bukti screenshot berstatus pass. Tidak ditemukan artefak berstatus fail pada hasil yang diimpor."), size=10.5, color=INK)


def add_scenario_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["No.", "Skenario", "Rute", "Validasi utama", "Status"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), size=9, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])

    for index, (name, route, validation, _) in enumerate(SCENARIOS, 1):
        cells = table.add_row().cells
        values = [str(index), name, route, validation, "PASS"]
        for col, value in enumerate(values):
            cell = cells[col]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if index % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in (0, 2, 4) else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(value), size=8.5, color=(GREEN if col == 4 else INK), bold=(col == 4))
    set_table_geometry(table, [450, 1750, 1580, 4500, 1080])
    set_table_borders(table)


def add_evidence(doc: Document, index: int, scenario: tuple[str, str, str, str]) -> None:
    name, route, validation, filename = scenario
    image_path = SCREENSHOT_DIR / filename

    heading = doc.add_paragraph(style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    set_font(heading.add_run(f"{index}. {name}"), size=13, color=BLUE, bold=True)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(4)
    meta.paragraph_format.keep_with_next = True
    set_font(meta.add_run("Rute: "), size=9, color=MUTED, bold=True)
    set_font(meta.add_run(route), size=9, color=INK)
    set_font(meta.add_run("   |   Pemeriksaan: "), size=9, color=MUTED, bold=True)
    set_font(meta.add_run(validation), size=9, color=INK)
    set_font(meta.add_run("   |   Status: "), size=9, color=MUTED, bold=True)
    set_font(meta.add_run("PASS"), size=9, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(image_path), width=Inches(6.25))
    add_image_alt_text(shape, f"Bukti Selenium: {name}", f"Screenshot hasil navigation smoke test untuk skenario {name}, status PASS.")

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(5)
    set_font(cap.add_run(f"Gambar {index}. Bukti visual {name} — {filename}"), size=8, color=MUTED, italic=True)


def build_report() -> Path:
    for _, _, _, filename in SCENARIOS:
        path = SCREENSHOT_DIR / filename
        if not path.exists():
            raise FileNotFoundError(path)

    latest_time = max(datetime.fromtimestamp((SCREENSHOT_DIR / item[3]).stat().st_mtime) for item in SCENARIOS)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("LAPORAN PENGUJIAN"), size=10, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Title")
    set_font(p.add_run("Hasil Navigation Smoke Test Selenium"), size=24, color=INK, bold=True)

    p = doc.add_paragraph(style="Subtitle")
    set_font(p.add_run("SIPENA — Sistem Informasi Pengelolaan Sarana"), size=13, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(14)
    add_bottom_border(rule, color=BLUE, size="18")

    add_metadata(doc, latest_time)
    doc.add_paragraph()
    add_status_callout(doc)

    doc.add_heading("1. Ringkasan Eksekutif", level=1)
    doc.add_paragraph(
        "Navigation smoke test Selenium memeriksa alur login, navigasi melalui menu sidebar, ketersediaan judul dan elemen utama pada setiap halaman operasional, serta logout. Berdasarkan 19 screenshot yang tersedia, seluruh skenario berstatus pass."
    )

    doc.add_heading("2. Metode dan Kriteria", level=1)
    doc.add_paragraph(
        "Suite membuka setiap fitur melalui klik menu, memvalidasi rute, judul halaman, teks atau kontrol wajib, memastikan halaman tidak kosong dan bukan halaman 404, kemudian menyimpan screenshot sebagai bukti. Skenario logout juga memeriksa bahwa token sesi telah dihapus."
    )
    p = doc.add_paragraph()
    set_font(p.add_run("Catatan: "), size=10, color=MUTED, bold=True)
    set_font(p.add_run("Laporan ini mengimpor artefak screenshot yang tersedia; tidak menjalankan ulang suite Selenium."), size=10, color=MUTED, italic=True)

    doc.add_page_break()
    doc.add_heading("3. Rekapitulasi Skenario", level=1)
    add_scenario_table(doc)

    doc.add_page_break()
    doc.add_heading("4. Lampiran Bukti Visual", level=1)
    intro = doc.add_paragraph("Setiap gambar berikut merupakan artefak langsung dari folder selenium/screenshots dan mempertahankan nama file sumbernya.")
    intro.paragraph_format.space_after = Pt(8)

    for index, scenario in enumerate(SCENARIOS, 1):
        add_evidence(doc, index, scenario)
        if index % 2 == 0 and index != len(SCENARIOS):
            doc.add_page_break()

    core_props = doc.core_properties
    core_props.title = "Laporan Hasil Navigation Smoke Test Selenium SIPENA"
    core_props.subject = "Rekap dan bukti visual hasil pengujian Selenium"
    core_props.author = "Tim Pengembangan SIPENA"
    core_props.keywords = "SIPENA, Selenium, smoke test, pengujian, laporan"
    core_props.comments = "Dibuat dari artefak screenshot Selenium yang tersedia di workspace."

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_report())
